import pytest
import docx
from pathlib import Path

from build_your_own_rag.models import SourceDocument, ExtractionStrategy
from build_your_own_rag.parsing.docx_parser import parse_docx

@pytest.fixture
def sample_docx_file(tmp_path):
    # Create a real docx file using python-docx
    file_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test DOCX file.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = 'Header 1'
    cell = table.cell(0, 1)
    cell.text = 'Header 2'
    cell = table.cell(1, 0)
    cell.text = 'Value 1'
    cell = table.cell(1, 1)
    cell.text = 'Value 2'
    
    doc.save(file_path)
    return file_path


def test_parse_docx_success(sample_docx_file):
    source = SourceDocument(
        source_id="test_id_docx",
        source_type="local_docx",
        path=sample_docx_file,
        filename="test.docx",
        extension="docx",
        size_bytes=sample_docx_file.stat().st_size,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    
    # Test with Docling (default/FAST)
    parsed = parse_docx(source, strategy=ExtractionStrategy.FAST)
    
    # Docling should succeed
    assert parsed.parser_name == "docling"
    assert parsed.source == source
    
    # It should export to markdown which includes the text
    assert "Test Document" in parsed.text
    assert "This is a test DOCX file." in parsed.text
    
    # Check that it extracted the table (in Docling markdown format, usually involves | and -)
    assert "Header 1" in parsed.text
    assert "Value 2" in parsed.text
    
    # Note: docling parser stores docx_metadata as empty dict if parsed by docling, 
    # since we only populate python-docx specific metadata during fallback.
    assert "docx" in parsed.metadata.format_specific

def test_parse_docx_fallback(monkeypatch, sample_docx_file):
    source = SourceDocument(
        source_id="test_id_docx_fallback",
        source_type="local_docx",
        path=sample_docx_file,
        filename="test.docx",
        extension="docx",
        size_bytes=sample_docx_file.stat().st_size,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    
    # Force Docling to fail to trigger fallback
    def mock_parse_with_docling(*args, **kwargs):
        raise RuntimeError("Mock Docling failure")
    
    import build_your_own_rag.parsing.docx_parser as docx_module
    monkeypatch.setattr(docx_module, "_parse_with_docling", mock_parse_with_docling)
    
    parsed = parse_docx(source, strategy=ExtractionStrategy.FAST)
    
    # Should be python_docx_fallback
    assert parsed.parser_name == "python_docx_fallback"
    assert parsed.status == "fallback"
    
    assert "Test Document" in parsed.text
    assert "This is a test DOCX file." in parsed.text
    
    # Should extract table contents as well
    assert "Header 1" in parsed.text
    assert "Value 2" in parsed.text
    
    # Python-docx metadata should be populated
    docx_meta = parsed.metadata.format_specific["docx"]
    assert "paragraph_count" in docx_meta
    assert docx_meta["table_count"] == 1
